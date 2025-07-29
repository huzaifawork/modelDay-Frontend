#!/usr/bin/env python3
"""
Create sample DOCX files for testing file upload functionality
Requires: pip install python-docx
"""

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    def create_modeling_contract():
        doc = Document()
        
        # Title
        title = doc.add_heading('MODELING CONTRACT AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contract details
        doc.add_heading('Contract Details', level=1)
        
        p = doc.add_paragraph()
        p.add_run('Client: ').bold = True
        p.add_run('Fashion Forward Magazine\n')
        p.add_run('Model: ').bold = True
        p.add_run('Sarah Johnson\n')
        p.add_run('Agency: ').bold = True
        p.add_run('Elite Modeling Agency\n')
        p.add_run('Date: ').bold = True
        p.add_run('January 15, 2025\n')
        p.add_run('Location: ').bold = True
        p.add_run('New York Studio, 123 Fashion Ave\n')
        
        # Terms
        doc.add_heading('Terms and Conditions', level=1)
        
        terms = [
            'Shoot Duration: 8 hours (9:00 AM - 5:00 PM)',
            'Rate: $500 per hour',
            'Total Fee: $4,000',
            'Usage Rights: Editorial use only, 1 year license',
            'Wardrobe: Provided by client',
            'Hair & Makeup: Professional team provided',
            'Payment Terms: Net 30 days',
            'Cancellation: 48 hours notice required'
        ]
        
        for term in terms:
            doc.add_paragraph(term, style='List Bullet')
        
        # Additional clauses
        doc.add_heading('Additional Clauses', level=1)
        
        doc.add_paragraph(
            'The model agrees to arrive on time and maintain professional conduct '
            'throughout the shoot. The client agrees to provide a safe working '
            'environment and all necessary equipment.'
        )
        
        doc.add_paragraph(
            'This contract is governed by the laws of New York State. Any disputes '
            'will be resolved through arbitration.'
        )
        
        # Signatures
        doc.add_heading('Signatures', level=1)
        
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.style = 'Table Grid'
        
        sig_table.cell(0, 0).text = 'Model Signature:'
        sig_table.cell(0, 1).text = 'Date:'
        sig_table.cell(1, 0).text = 'Client Signature:'
        sig_table.cell(1, 1).text = 'Date:'
        sig_table.cell(2, 0).text = 'Agent Signature:'
        sig_table.cell(2, 1).text = 'Date:'
        
        filename = 'modeling_contract_sample.docx'
        doc.save(filename)
        return filename
    
    def create_model_portfolio():
        doc = Document()
        
        # Title
        title = doc.add_heading('MODEL PORTFOLIO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Model info
        doc.add_heading('Model Information', level=1)
        
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        model_info = [
            ('Name:', 'Alexandra Smith'),
            ('Age:', '22'),
            ('Height:', '5\'9" (175 cm)'),
            ('Measurements:', '34-24-36'),
            ('Hair Color:', 'Blonde'),
            ('Eye Color:', 'Blue'),
            ('Agency:', 'Premier Talent Group'),
            ('Experience:', '3 years professional modeling')
        ]
        
        for i, (label, value) in enumerate(model_info):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        # Experience
        doc.add_heading('Professional Experience', level=1)
        
        experiences = [
            'Vogue Magazine - Editorial Spread (2024)',
            'Calvin Klein - Commercial Campaign (2023)',
            'New York Fashion Week - Runway Shows (2023-2024)',
            'H&M - Print Advertisement (2023)',
            'Local Fashion Boutiques - Catalog Shoots (2022-2024)'
        ]
        
        for exp in experiences:
            doc.add_paragraph(exp, style='List Bullet')
        
        # Skills
        doc.add_heading('Skills & Specialties', level=1)
        
        skills = [
            'Editorial Photography',
            'Commercial Modeling',
            'Runway Walking',
            'Product Photography',
            'Fitness Modeling',
            'Beauty Shots'
        ]
        
        for skill in skills:
            doc.add_paragraph(skill, style='List Bullet')
        
        # Contact
        doc.add_heading('Contact Information', level=1)
        
        contact = doc.add_paragraph()
        contact.add_run('Email: ').bold = True
        contact.add_run('alexandra.smith@premiertalent.com\n')
        contact.add_run('Phone: ').bold = True
        contact.add_run('+1 (555) 987-6543\n')
        contact.add_run('Agent: ').bold = True
        contact.add_run('Michael Chen - m.chen@premiertalent.com\n')
        contact.add_run('Portfolio: ').bold = True
        contact.add_run('www.alexandrasmith-model.com')
        
        filename = 'model_portfolio_sample.docx'
        doc.save(filename)
        return filename
    
    def create_booking_details():
        doc = Document()
        
        # Title
        title = doc.add_heading('BOOKING DETAILS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Booking info
        doc.add_heading('Booking Information', level=1)
        
        booking_table = doc.add_table(rows=10, cols=2)
        booking_table.style = 'Table Grid'
        
        booking_info = [
            ('Booking ID:', 'BK-2025-0115-001'),
            ('Client:', 'Luxury Fashion Brand'),
            ('Project:', 'Spring Collection Catalog'),
            ('Date:', 'January 20, 2025'),
            ('Time:', '10:00 AM - 6:00 PM'),
            ('Location:', 'Downtown Studio, Los Angeles'),
            ('Model:', 'Jessica Williams'),
            ('Rate:', '$800/day'),
            ('Usage:', 'Print & Digital Catalog'),
            ('Status:', 'Confirmed')
        ]
        
        for i, (label, value) in enumerate(booking_info):
            booking_table.cell(i, 0).text = label
            booking_table.cell(i, 1).text = value
        
        # Requirements
        doc.add_heading('Requirements', level=1)
        
        requirements = [
            'Bring natural makeup and hair styling tools',
            'Wear comfortable, form-fitting undergarments',
            'Arrive 30 minutes early for preparation',
            'Professional attitude and punctuality required',
            'No food or drinks near wardrobe area'
        ]
        
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # Notes
        doc.add_heading('Additional Notes', level=1)
        
        doc.add_paragraph(
            'This is a high-profile catalog shoot for a luxury brand. '
            'Professional conduct is essential. Parking is available '
            'in the building garage. Lunch will be provided.'
        )
        
        filename = 'booking_details_sample.docx'
        doc.save(filename)
        return filename
    
    # Create all sample files
    print("📄 Creating sample DOCX files...")
    
    files_created = []
    files_created.append(create_modeling_contract())
    files_created.append(create_model_portfolio())
    files_created.append(create_booking_details())
    
    print(f"\n✅ Created {len(files_created)} sample DOCX files:")
    for i, filename in enumerate(files_created, 1):
        print(f"{i}. {filename}")
    
    print(f"\n🧪 How to test:")
    print(f"1. Go to http://localhost:3000/#/new-option")
    print(f"2. Click 'Add Files' button")
    print(f"3. Select one of the created DOCX files")
    print(f"4. Check if file preview appears")
    print(f"5. Submit the form to test file upload")

except ImportError:
    print("❌ python-docx not installed. Install with: pip install python-docx")
except Exception as e:
    print(f"❌ Error creating DOCX files: {e}")
